"""
Documentation generator for Counter-UAS Radar Tracker project.
Produces a detailed DOCX file.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), kwargs.get(edge, 'single'))
        tag.set(qn('w:sz'), kwargs.get('sz', '4'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', '2E4057'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    run = h.runs[0] if h.runs else h.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)  # deep navy
    elif level == 2:
        run.font.color.rgb = RGBColor(0x1A, 0x6B, 0x3C)  # deep green
    elif level == 3:
        run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)
    return h

def add_para(doc, text, bold=False, italic=False, size=10.5, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def add_code_block(doc, code_text):
    """Add a shaded monospace block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F4F8')
    pPr.append(shd)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)
    return p

def add_table(doc, headers, rows, col_widths=None):
    """Add a styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_bg(cell, '1B3A6B')
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        set_cell_border(cell, color='FFFFFF')

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = 'EEF2F7' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, fill)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
            cell.paragraphs[0].paragraph_format.space_before = Pt(1)
            cell.paragraphs[0].paragraph_format.space_after = Pt(1)
            set_cell_border(cell, color='C8D6E5')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table

def page_break(doc):
    doc.add_page_break()

# ---------------------------------------------------------------------------
# Document build
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # =========================================================================
    # COVER PAGE
    # =========================================================================
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Counter-UAS Radar Tracker")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("System Architecture & Technical Documentation")
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(0x1A, 0x6B, 0x3C)

    doc.add_paragraph()

    ver_p = doc.add_paragraph()
    ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver_run = ver_p.add_run("Version 1.0.0  |  February 2026")
    ver_run.font.size = Pt(11)
    ver_run.font.color.rgb = RGBColor(0x55, 0x66, 0x77)

    doc.add_paragraph()
    doc.add_paragraph()

    proj_p = doc.add_paragraph()
    proj_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    proj_run = proj_p.add_run("Zoppler Projects  –  Tracker_CUxS")
    proj_run.font.size = Pt(10)
    proj_run.italic = True
    proj_run.font.color.rgb = RGBColor(0x88, 0x99, 0xAA)

    page_break(doc)

    # =========================================================================
    # 1. EXECUTIVE SUMMARY
    # =========================================================================
    add_heading(doc, "1. Executive Summary", 1)
    add_para(doc, (
        "The Counter-UAS Radar Tracker (Tracker_CUxS) is a C++17 real-time signal processing "
        "system designed to detect, track, and classify Unmanned Aerial Systems (UAS) using radar "
        "sensor data. It implements a full multi-target tracking pipeline — from raw radar detections "
        "through preprocessing, clustering, Kalman-based multi-model prediction, data association, "
        "and track management — outputting track estimates over a UDP network interface for downstream "
        "consumers such as command-and-control systems or display terminals."
    ))
    add_para(doc, (
        "The system is built around three interchangeable algorithmic strategies for clustering "
        "(DBSCAN, Range-based, Range-Strength-based) and three for data association (Mahalanobis "
        "greedy, Global Nearest Neighbour / GNN, and Joint Probabilistic Data Association / JPDA), "
        "all selectable at runtime via a JSON configuration file without recompilation. Prediction "
        "is handled by a five-model Interacting Multiple Model (IMM) filter covering Constant Velocity "
        "(CV), two Constant Acceleration (CA) variants, and two Coordinated Turn Rate (CTR) models."
    ))
    add_para(doc, (
        "Alongside the core tracker, the package includes a Qt5 GUI display module, a synthetic DSP "
        "target injector, a console display simulator, and a binary log extractor/replayer — providing "
        "a self-contained development and test environment."
    ))

    # =========================================================================
    # 2. SYSTEM OVERVIEW
    # =========================================================================
    add_heading(doc, "2. System Overview", 1)

    add_heading(doc, "2.1 High-Level Architecture", 2)
    add_para(doc, (
        "The tracker is structured as a pipeline of loosely coupled processing stages, each "
        "implemented as an independent library and orchestrated by a central pipeline controller. "
        "Data flows from the UDP receiver through sequential processing stages and exits via the "
        "track sender."
    ))
    add_code_block(doc,
        "  [Radar Sensor / DSP Injector]\n"
        "          |\n"
        "          v  (UDP — SPDetectionMessage)\n"
        "  ┌───────────────────┐\n"
        "  │  Detection        │  DetectionReceiver (threaded)\n"
        "  │  Receiver         │\n"
        "  └────────┬──────────┘\n"
        "           │  queue\n"
        "           v\n"
        "  ┌────────────────────────────────────────────────────────┐\n"
        "  │                  Tracker Pipeline                      │\n"
        "  │                                                        │\n"
        "  │  Preprocessing  →  Clustering  →  TrackManager        │\n"
        "  │                                      │                 │\n"
        "  │                           ┌──────────┼──────────┐     │\n"
        "  │                       Prediction  Assoc.  Initiation  │\n"
        "  │                                                        │\n"
        "  └────────────────────────────┬───────────────────────────┘\n"
        "                               │\n"
        "                               v  (UDP — TrackTableMessage)\n"
        "                      [Display Module / C2 System]\n"
    )

    add_heading(doc, "2.2 Key Design Principles", 2)
    for item in [
        ("Strategy Pattern", "Clustering, prediction, and association are selectable algorithms with a common interface, chosen at startup from the JSON config."),
        ("Plugin / Factory Instantiation", "ClusterEngine, AssociationEngine, and TrackerPipeline act as factories that instantiate and own the selected algorithm object."),
        ("Lock-Free-Style Threading", "The pipeline uses a producer-consumer queue with a condition variable between the receiver thread and the processor thread, minimising blocking."),
        ("Binary Wire Protocol", "Custom little-endian binary serialisation (no third-party protocol library) enables compact, low-latency UDP transmission."),
        ("Configurable Lifecycle", "Track initiation, confirmation, coasting, and deletion are all governed by JSON-configurable M-of-N counters, quality thresholds, and age limits."),
        ("Full Audit Trail", "A structured binary log records every stage of the pipeline (raw, preprocessed, clustered, predicted, associated, initiated, updated, deleted, sent) for post-mission analysis and replay."),
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(item[0] + ": ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(item[1])
        r2.font.size = Pt(10)

    # =========================================================================
    # 3. REPOSITORY STRUCTURE
    # =========================================================================
    add_heading(doc, "3. Repository Structure", 1)
    add_table(doc,
        ["Path", "Description"],
        [
            ["CMakeLists.txt",             "Top-level CMake build script defining all libraries and executables"],
            ["config/tracker_config.json", "Runtime configuration file (algorithm selection, thresholds, network ports)"],
            ["idl/messages.idl",           "Interface definition language file describing the binary wire messages"],
            ["include/common/",            "Shared headers: types, constants, config structs, logger, matrix ops, UDP socket"],
            ["include/receiver/",          "DetectionReceiver interface header"],
            ["include/preprocessing/",     "Preprocessor header"],
            ["include/clustering/",        "ClusterEngine + three clusterer headers"],
            ["include/prediction/",        "IMotionModel interface + CV/CA/CTR/IMM headers"],
            ["include/association/",       "AssociationEngine + three associator headers"],
            ["include/track_management/",  "Track, TrackInitiator, TrackManager headers"],
            ["include/sender/",            "TrackSender header"],
            ["include/pipeline/",          "TrackerPipeline header"],
            ["src/",                       "All .cpp implementation files, mirroring the include/ structure"],
            ["src/main.cpp",               "Executable entry point with signal handling and config path resolution"],
            ["simulators/dsp_injector/",   "Synthetic multi-target radar detection generator"],
            ["simulators/display_module/", "Console track display consumer"],
            ["simulators/log_extractor/",  "Binary log parser, CSV exporter, and replay injector"],
            ["qt_display_module/",         "Qt5 GUI track display application"],
            ["build/",                     "Out-of-source CMake build directory (Visual Studio / MSBuild)"],
        ],
        col_widths=[2.2, 4.4]
    )

    page_break(doc)

    # =========================================================================
    # 4. BUILD SYSTEM
    # =========================================================================
    add_heading(doc, "4. Build System", 1)

    add_heading(doc, "4.1 CMake Configuration", 2)
    add_para(doc, (
        "The project uses CMake 3.14+ with C++17. On Windows (MSVC) it enables /W4 warning level "
        "and links ws2_32 for Winsock; on Linux/macOS it uses -Wall -Wextra -O2 and links pthreads. "
        "The build produces nine static libraries and four executables."
    ))

    add_heading(doc, "4.2 Static Libraries", 2)
    add_table(doc,
        ["Library", "Source Files", "Dependencies"],
        [
            ["cuas_common",           "config.cpp, logger.cpp, udp_socket.cpp",                              "ws2_32 / pthread"],
            ["cuas_receiver",         "detection_receiver.cpp",                                              "cuas_common"],
            ["cuas_preprocessing",    "preprocessor.cpp",                                                    "cuas_common"],
            ["cuas_clustering",       "cluster_engine.cpp, dbscan_clusterer.cpp, range_clusterer.cpp,\nrange_strength_clusterer.cpp", "cuas_common"],
            ["cuas_prediction",       "cv_model.cpp, ca_model.cpp, ctr_model.cpp, imm_filter.cpp",          "cuas_common"],
            ["cuas_association",      "association_engine.cpp, mahalanobis_associator.cpp,\ngnn_associator.cpp, jpda_associator.cpp", "cuas_common, cuas_prediction"],
            ["cuas_track_management", "track.cpp, track_initiator.cpp, track_manager.cpp",                  "cuas_common, cuas_preprocessing, cuas_clustering, cuas_prediction, cuas_association"],
            ["cuas_sender",           "track_sender.cpp",                                                    "cuas_common"],
            ["cuas_pipeline",         "tracker_pipeline.cpp",                                                "cuas_common, cuas_receiver, cuas_track_management, cuas_sender"],
        ],
        col_widths=[1.8, 2.8, 2.0]
    )

    add_heading(doc, "4.3 Executables", 2)
    add_table(doc,
        ["Executable",      "Entry Point",                         "Links Against"],
        [
            ["cuas_tracker",    "src/main.cpp",                        "cuas_pipeline (which transitively pulls all libs)"],
            ["dsp_injector",    "simulators/dsp_injector/dsp_injector.cpp",     "cuas_common"],
            ["display_module",  "simulators/display_module/display_module.cpp", "cuas_common"],
            ["log_extractor",   "simulators/log_extractor/log_extractor.cpp",   "cuas_common"],
        ],
        col_widths=[1.6, 2.6, 2.4]
    )

    add_heading(doc, "4.4 Build Commands", 2)
    add_para(doc, "Configure (first time, from project root):")
    add_code_block(doc, "cmake -B build -S .")
    add_para(doc, "Build Debug configuration:")
    add_code_block(doc, "cmake --build build --config Debug")
    add_para(doc, "Build Release configuration:")
    add_code_block(doc, "cmake --build build --config Release")
    add_para(doc, (
        "After a successful build the output binaries and a copy of tracker_config.json are placed "
        "in build/Debug/ or build/Release/ depending on the selected configuration."
    ))

    page_break(doc)

    # =========================================================================
    # 5. CONFIGURATION
    # =========================================================================
    add_heading(doc, "5. Configuration Reference", 1)
    add_para(doc, (
        "All runtime parameters are read from config/tracker_config.json at startup. "
        "The tracker searches for the file first in the executable's own directory, then up two "
        "parent levels, allowing the same binary to be used from different working directories."
    ))

    add_heading(doc, "5.1 System Parameters", 2)
    add_table(doc,
        ["Parameter", "Default", "Description"],
        [
            ["cycle_period_ms",        "100",   "Processing cycle period in milliseconds (10 Hz)"],
            ["max_detections_per_scan","256",   "Maximum detections accepted per radar scan"],
            ["max_tracks",             "64",    "Maximum concurrent tracks maintained"],
            ["enable_logging",         "true",  "Enable/disable binary file logging"],
            ["log_file",               "tracker_log.bin", "Output log file path"],
            ["log_level",              "INFO",  "Logging verbosity (DEBUG / INFO / WARN / ERROR)"],
        ],
        col_widths=[2.2, 1.2, 3.2]
    )

    add_heading(doc, "5.2 Network Parameters", 2)
    add_table(doc,
        ["Parameter", "Default", "Description"],
        [
            ["listen_port",     "5000",      "UDP port on which tracker receives radar detections"],
            ["send_ip",         "127.0.0.1", "Destination IP for track output messages"],
            ["send_port",       "5001",      "Destination UDP port for track output"],
            ["recv_buffer_size","65536",     "UDP receive socket buffer size (bytes)"],
            ["send_buffer_size","65536",     "UDP send socket buffer size (bytes)"],
        ],
        col_widths=[2.0, 1.4, 3.2]
    )

    add_heading(doc, "5.3 Preprocessing Parameters", 2)
    add_table(doc,
        ["Parameter",        "Default", "Description"],
        [
            ["min_range_m",    "30",    "Minimum valid detection range (m)"],
            ["max_range_m",    "20000", "Maximum valid detection range (m)"],
            ["min_azimuth_deg","-180",  "Minimum azimuth angle (degrees)"],
            ["max_azimuth_deg","180",   "Maximum azimuth angle (degrees)"],
            ["min_elevation_deg","-10", "Minimum elevation angle (degrees)"],
            ["max_elevation_deg","90",  "Maximum elevation angle (degrees)"],
            ["min_snr_db",     "5.0",   "Minimum Signal-to-Noise Ratio threshold (dB)"],
            ["min_rcs_dbsm",   "-20.0", "Minimum Radar Cross-Section threshold (dBsm)"],
            ["min_strength",   "0.1",   "Minimum signal strength (normalised)"],
        ],
        col_widths=[2.0, 1.2, 3.4]
    )

    add_heading(doc, "5.4 Clustering Parameters", 2)
    add_para(doc, "The clustering method is selected via the \"method\" field: \"DBSCAN\", \"RangeBased\", or \"RangeStrength\".")
    add_table(doc,
        ["Parameter",             "Applies To",           "Default", "Description"],
        [
            ["method",                "all",                  "DBSCAN",  "Clustering algorithm selection"],
            ["epsilon",               "DBSCAN",               "150.0",   "Neighbourhood radius in composite distance metric"],
            ["min_points",            "DBSCAN",               "1",       "Minimum points to form a core point"],
            ["range_weight",          "DBSCAN",               "1.0",     "Weight for range component in DBSCAN metric"],
            ["azimuth_weight",        "DBSCAN",               "57.3",    "Weight for azimuth component (converts rad to m-equivalent)"],
            ["elevation_weight",      "DBSCAN",               "57.3",    "Weight for elevation component"],
            ["range_gate_m",          "RangeBased/RS",        "200.0",   "Gate size in range dimension (m)"],
            ["strength_gate",         "RangeStrength",        "0.5",     "Gate size in signal strength dimension"],
        ],
        col_widths=[2.0, 1.6, 1.0, 2.0]
    )

    add_heading(doc, "5.5 Prediction / IMM Parameters", 2)
    add_para(doc, (
        "The IMM filter runs five motion models simultaneously. Each model has its own process "
        "noise standard deviations. The model transition matrix (Markov chain) governs how "
        "probability mass moves between models each cycle."
    ))
    add_table(doc,
        ["Model",  "Index", "Process Noise σ (m/s² or rad/s)", "Purpose"],
        [
            ["CV",   "0", "accel_std: 2.0",       "Constant Velocity — straight, level flight"],
            ["CA1",  "1", "accel_std: 5.0",        "Constant Acceleration (low) — gentle manoeuvre"],
            ["CA2",  "2", "accel_std: 15.0",       "Constant Acceleration (high) — aggressive manoeuvre"],
            ["CTR1", "3", "accel_std: 3.0, turn_rate_std: 0.05", "Coordinated Turn (slow) — gentle banking"],
            ["CTR2", "4", "accel_std: 5.0, turn_rate_std: 0.15", "Coordinated Turn (fast) — tight banking"],
        ],
        col_widths=[0.8, 0.7, 2.5, 2.6]
    )

    add_heading(doc, "5.6 Association Parameters", 2)
    add_para(doc, "The association method is selected via the \"method\" field: \"Mahalanobis\", \"GNN\", or \"JPDA\".")
    add_table(doc,
        ["Parameter",              "Default", "Description"],
        [
            ["method",             "GNN",     "Data association algorithm"],
            ["gate_threshold",     "16.0",    "Chi-squared gate (df=3): 16.0 ≈ 99.7% for 3-DoF Gaussian"],
            ["max_association_dist","500.0",  "Hard range limit on candidate pairs (m, Euclidean)"],
        ],
        col_widths=[2.2, 1.2, 3.2]
    )

    add_heading(doc, "5.7 Track Management Parameters", 2)
    add_table(doc,
        ["Parameter",             "Default", "Description"],
        [
            ["m_of_n_m",           "3",     "Hits required for track confirmation (M)"],
            ["m_of_n_n",           "5",     "Observation window for confirmation (N)"],
            ["max_coasting_cycles","5",     "Maximum consecutive misses before track deletion"],
            ["min_track_quality",  "0.1",   "Quality floor; track deleted if quality falls below this"],
            ["quality_decay_rate", "0.15",  "Quality decrement per missed scan"],
            ["quality_boost_rate", "0.05",  "Quality increment per successful association"],
            ["max_track_age",      "3600",  "Maximum track lifetime in scans"],
            ["deletion_min_hits",  "3",     "Minimum lifetime hits; track deleted if not reached before max coasting"],
        ],
        col_widths=[2.2, 1.0, 3.4]
    )

    page_break(doc)

    # =========================================================================
    # 6. DATA STRUCTURES & WIRE PROTOCOL
    # =========================================================================
    add_heading(doc, "6. Data Structures and Wire Protocol", 1)

    add_heading(doc, "6.1 SPDetectionMessage  (Radar → Tracker)", 2)
    add_para(doc, (
        "Sent by the radar's Digital Signal Processor (or the dsp_injector simulator) to the "
        "tracker's listen port. Each message carries up to max_detections detections for a single "
        "radar scan."
    ))
    add_table(doc,
        ["Field",           "Type",    "Unit",   "Description"],
        [
            ["msg_id",          "uint16",  "—",      "Message identifier: 0x0001"],
            ["scan_id",         "uint32",  "—",      "Monotonically increasing scan counter"],
            ["timestamp_ms",    "uint64",  "ms",     "Sensor epoch timestamp"],
            ["num_detections",  "uint16",  "—",      "Number of valid detection entries"],
            ["detections[]",    "struct",  "—",      "Array of Detection structs (see below)"],
        ],
        col_widths=[1.8, 1.0, 0.8, 3.0]
    )
    add_heading(doc, "Detection struct fields:", 3)
    add_table(doc,
        ["Field",        "Type",   "Unit",   "Description"],
        [
            ["range_m",      "float32","m",      "Slant range to detection"],
            ["azimuth_rad",  "float32","rad",    "Azimuth angle (positive = East)"],
            ["elevation_rad","float32","rad",    "Elevation angle (positive = up)"],
            ["strength",     "float32","—",      "Normalised signal amplitude [0, 1]"],
            ["snr_db",       "float32","dB",     "Signal-to-Noise Ratio"],
            ["rcs_dbsm",     "float32","dBsm",   "Radar Cross-Section estimate"],
            ["micro_doppler","float32","Hz",     "Micro-Doppler signature (blade flash)"],
        ],
        col_widths=[1.8, 1.0, 0.8, 3.0]
    )

    add_heading(doc, "6.2 TrackUpdateMessage  (Tracker → Display)", 2)
    add_para(doc, "Sent for each active track every output cycle. Message ID 0x0002.")
    add_table(doc,
        ["Field",            "Type",    "Unit",   "Description"],
        [
            ["msg_id",           "uint16",  "—",      "0x0002"],
            ["track_id",         "uint32",  "—",      "Unique track identifier"],
            ["timestamp_ms",     "uint64",  "ms",     "Timestamp of last update"],
            ["status",           "uint8",   "—",      "0=Tentative, 1=Confirmed, 2=Coasting, 3=Deleted"],
            ["classification",   "uint8",   "—",      "0=Unknown, 1=FixedWing, 2=Rotor, 3=Micro, 4=Balloon"],
            ["range_m",          "float32", "m",      "Estimated slant range"],
            ["azimuth_rad",      "float32", "rad",    "Estimated azimuth"],
            ["elevation_rad",    "float32", "rad",    "Estimated elevation"],
            ["range_rate_m_s",   "float32", "m/s",    "Radial velocity (positive = moving away)"],
            ["pos_x_m",          "float32", "m",      "Cartesian East position"],
            ["pos_y_m",          "float32", "m",      "Cartesian North position"],
            ["pos_z_m",          "float32", "m",      "Cartesian Up position"],
            ["vel_x_m_s",        "float32", "m/s",    "Cartesian East velocity"],
            ["vel_y_m_s",        "float32", "m/s",    "Cartesian North velocity"],
            ["vel_z_m_s",        "float32", "m/s",    "Cartesian Up velocity"],
            ["quality",          "float32", "—",      "Track quality metric [0.0, 1.0]"],
            ["hits",             "uint32",  "—",      "Cumulative successful associations"],
            ["misses",           "uint32",  "—",      "Consecutive missed associations"],
            ["age",              "uint32",  "—",      "Track age in scans"],
        ],
        col_widths=[1.8, 1.0, 0.8, 3.0]
    )

    add_heading(doc, "6.3 TrackTableMessage  (Tracker → Display)", 2)
    add_para(doc, (
        "A batch wrapper sent once per output cycle containing all active tracks. Message ID 0x0003. "
        "Consists of a header with track_count followed by track_count TrackUpdateMessage payloads "
        "packed sequentially."
    ))

    page_break(doc)

    # =========================================================================
    # 7. PROCESSING PIPELINE
    # =========================================================================
    add_heading(doc, "7. Processing Pipeline", 1)

    add_heading(doc, "7.1 Detection Receiver", 2)
    add_para(doc, (
        "DetectionReceiver runs in a dedicated thread. It blocks on a UDP recvfrom() call, "
        "deserialises the incoming binary payload into an SPDetectionMessage, and invokes a "
        "registered callback. The callback (implemented in TrackerPipeline) pushes the message "
        "onto a thread-safe queue and signals the processor thread via a condition variable."
    ))

    add_heading(doc, "7.2 Preprocessing", 2)
    add_para(doc, (
        "The Preprocessor applies six independent gates to each Detection in a scan, discarding "
        "detections that fail any gate. Gates are applied in this order:"
    ))
    for g in ["Range gate [min_range_m, max_range_m]",
              "Azimuth gate [min_azimuth_deg, max_azimuth_deg] (converted to radians internally)",
              "Elevation gate [min_elevation_deg, max_elevation_deg]",
              "SNR gate [min_snr_db, ∞)",
              "RCS gate [min_rcs_dbsm, ∞)",
              "Strength gate [min_strength, ∞)"]:
        add_bullet(doc, g)

    add_heading(doc, "7.3 Clustering", 2)
    add_para(doc, (
        "Clustering groups spatially close detections into Cluster objects, each containing a "
        "weighted centroid (range, azimuth, elevation, strength, SNR, RCS) and a list of constituent "
        "detections. Three algorithms are available:"
    ))

    add_heading(doc, "DBSCAN (Density-Based Spatial Clustering of Applications with Noise)", 3)
    add_para(doc, (
        "Assigns each detection a core-point status based on an epsilon-normalised composite distance "
        "metric combining range, azimuth, and elevation differences weighted by their respective "
        "configuration weights. A single-point cluster is created for noise detections rather than "
        "discarding them, ensuring no detections are lost. Cluster centroids are computed as "
        "arithmetic means of member detections."
    ))
    add_code_block(doc,
        "distance(a, b) = sqrt(\n"
        "    (range_weight  * Δrange)²  +\n"
        "    (azimuth_weight * Δazimuth)² +\n"
        "    (elevation_weight * Δelevation)²\n"
        ")\n"
        "Core point: |N_ε(p)| >= min_points"
    )

    add_heading(doc, "Range-Based Clustering", 3)
    add_para(doc, (
        "Sorts detections by range and performs greedy sequential grouping. A detection is added to "
        "the current cluster if its range is within range_gate_m of the cluster's running centroid; "
        "otherwise a new cluster is started."
    ))

    add_heading(doc, "Range-Strength Clustering", 3)
    add_para(doc, (
        "Extends Range-Based clustering with an additional strength gate. Both the range gap and the "
        "strength difference must fall within their respective thresholds for a detection to join a cluster."
    ))

    add_heading(doc, "7.4 Prediction (IMM Filter)", 2)
    add_para(doc, (
        "The Interacting Multiple Model (IMM) filter maintains a bank of five Kalman filters, one "
        "per motion model. At each cycle it:"
    ))
    for step, desc in [
        ("1. Interaction / Mixing", "Computes mixed initial conditions for each model from the "
            "previous cycle's per-model state estimates, weighted by the model transition "
            "probability matrix and mode probabilities."),
        ("2. Mode-Conditioned Prediction", "Each model propagates its mixed state through its own "
            "state transition matrix F and adds process noise Q."),
        ("3. Kalman Update", "If a measurement is available, each model computes its innovation, "
            "innovation covariance S, Kalman gain K, and updated state/covariance. Likelihood "
            "Λᵢ is computed from the innovation PDF."),
        ("4. Mode Probability Update", "New mode probabilities μᵢ are proportional to Λᵢ * "
            "predicted mode probability (normalised to sum to 1.0)."),
        ("5. State Fusion", "The overall IMM estimate is the probability-weighted sum of the "
            "per-model updated states and covariances."),
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(step + ": ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(desc)
        r2.font.size = Pt(10)

    add_heading(doc, "State Vector", 3)
    add_para(doc, (
        "The state vector is 9-dimensional:  x = [px, py, pz, vx, vy, vz, ax, ay, az]  "
        "(position, velocity, acceleration in Cartesian ENU coordinates)."
    ))

    add_heading(doc, "Motion Models", 3)
    add_table(doc,
        ["Model", "F (state transition)",          "Q (process noise)"],
        [
            ["CV",   "p += v*dt; v = v; a = 0",     "Singer model: accel_std² on acceleration diagonal"],
            ["CA1/2","p += v*dt + ½a*dt²; v += a*dt; a = a*(1-decay)", "Full 9-state process noise with acceleration correlation"],
            ["CTR",  "Rotation of vx/vy by ω*dt; ω estimated from atan2(vy,vx) changes", "accel_std on v/a; turn_rate_std on ω"],
        ],
        col_widths=[0.8, 3.0, 2.8]
    )

    add_heading(doc, "7.5 Data Association", 2)
    add_para(doc, (
        "Data association pairs predicted track positions with incoming cluster centroids. "
        "Three algorithms are provided:"
    ))

    add_heading(doc, "Mahalanobis Greedy", 3)
    add_para(doc, (
        "For each track (sorted by quality descending), finds the cluster with the smallest "
        "Mahalanobis distance within the gate_threshold. Assigned clusters are removed from "
        "the candidate set. O(T × C) complexity."
    ))
    add_code_block(doc,
        "d_M²(z, z_pred) = (z - z_pred)ᵀ · S⁻¹ · (z - z_pred)\n"
        "where S = H·P·Hᵀ + R  (innovation covariance)"
    )

    add_heading(doc, "GNN (Global Nearest Neighbour)", 3)
    add_para(doc, (
        "Builds a cost matrix of Mahalanobis distances for all gated track-cluster pairs. "
        "Applies a simplified Hungarian-style row/column reduction to find the globally "
        "optimal one-to-one assignment. Preferred over greedy Mahalanobis when targets are "
        "closely spaced."
    ))

    add_heading(doc, "JPDA (Joint Probabilistic Data Association)", 3)
    add_para(doc, (
        "Computes association weights (β values) for all gated track-cluster pairs using "
        "the JPDA approximation. Each track's update is a weighted average over all its "
        "gated clusters. Preferred in dense clutter environments where multiple measurements "
        "could plausibly originate from the same target."
    ))

    add_heading(doc, "7.6 Track Management", 2)

    add_heading(doc, "Track Initiation (M-of-N)", 3)
    add_para(doc, (
        "Unassociated clusters are compared against tentative track candidates. If an unassociated "
        "cluster falls within a spatial gate of an existing candidate, a hit is recorded. A candidate "
        "with M or more hits within the last N scans is promoted to a confirmed Tentative track, "
        "with an initial state estimated from a two-point velocity approximation and a diagonal "
        "covariance matrix. Candidates older than N scans are purged."
    ))

    add_heading(doc, "Track Status Lifecycle", 3)
    add_code_block(doc,
        "  [Candidate]  -- M-of-N hits --> [Tentative]  -- association --> [Confirmed]\n"
        "                                       |                                |\n"
        "                                   misses                           misses\n"
        "                                       |                                |\n"
        "                                   [Deleted]                      [Coasting]\n"
        "                                                                        |\n"
        "                                                              max_coasting_cycles\n"
        "                                                                        |\n"
        "                                                                   [Deleted]"
    )

    add_heading(doc, "Track Quality", 3)
    add_para(doc, (
        "Each track maintains a quality metric in [0, 1]. On a hit, quality += quality_boost_rate "
        "(capped at 1.0). On a miss, quality -= quality_decay_rate (floored at 0.0). "
        "Tracks whose quality falls below min_track_quality are deleted."
    ))

    add_heading(doc, "Track Classification", 3)
    add_para(doc, (
        "A simple heuristic classifier operates after each update. It inspects the horizontal "
        "speed derived from the track's velocity state and the IMM mode probabilities:"
    ))
    add_table(doc,
        ["Classification",  "Criteria"],
        [
            ["FixedWing",       "Horizontal speed > 50 m/s AND CV or CA model dominates (prob > 0.4)"],
            ["Rotary",          "Horizontal speed 5–50 m/s AND CTR model has moderate probability (> 0.3)"],
            ["Micro",           "Horizontal speed < 10 m/s"],
            ["Balloon",         "Horizontal speed < 5 m/s AND vertical speed > horizontal speed"],
            ["Unknown",         "Insufficient data or no clear model dominance"],
        ],
        col_widths=[1.5, 5.1]
    )

    add_heading(doc, "7.7 Track Sender", 2)
    add_para(doc, (
        "TrackSender packages all non-deleted tracks into a TrackTableMessage and transmits it "
        "via UDP to the configured send_ip:send_port. Optionally, tracks marked Deleted can be "
        "included in one final transmission before being removed from the table."
    ))

    page_break(doc)

    # =========================================================================
    # 8. COMPONENT REFERENCE
    # =========================================================================
    add_heading(doc, "8. Component Reference", 1)

    add_heading(doc, "8.1 Common Library  (cuas_common)", 2)

    add_heading(doc, "types.h", 3)
    add_para(doc, "Defines all shared data structures used across the pipeline:")
    add_table(doc,
        ["Type",                "Kind",   "Purpose"],
        [
            ["Detection",           "struct", "Single radar detection in spherical coordinates with quality metrics"],
            ["SPDetectionMessage",  "struct", "UDP payload: array of detections for one radar scan"],
            ["Cluster",             "struct", "Aggregated cluster centroid + member detections"],
            ["StateVector",         "struct", "9-element ENU position/velocity/acceleration vector"],
            ["CovarianceMatrix",    "struct", "9×9 symmetric positive-definite covariance"],
            ["MeasurementVector",   "struct", "3-element [range, azimuth, elevation] measurement"],
            ["MeasurementCovMatrix","struct", "3×3 measurement noise covariance"],
            ["TrackUpdateMessage",  "struct", "Serialised track estimate for wire transmission"],
            ["TrackTableMessage",   "struct", "Batch wrapper for multiple TrackUpdateMessage entries"],
            ["TrackStatus",         "enum",   "Tentative / Confirmed / Coasting / Deleted"],
            ["Classification",      "enum",   "Unknown / FixedWing / Rotary / Micro / Balloon"],
            ["LogRecordType",       "enum",   "Pipeline stage identifiers for binary log records"],
        ],
        col_widths=[2.0, 0.8, 3.8]
    )

    add_heading(doc, "constants.h", 3)
    add_table(doc,
        ["Constant",          "Value",        "Description"],
        [
            ["PI",                "3.14159…",     "Mathematical π"],
            ["RAD2DEG",           "180/π",        "Radians-to-degrees conversion factor"],
            ["DEG2RAD",           "π/180",        "Degrees-to-radians conversion factor"],
            ["MSG_ID_DETECTION",  "0x0001",       "Wire protocol: radar detection message"],
            ["MSG_ID_TRACK",      "0x0002",       "Wire protocol: single track update"],
            ["MSG_ID_TRACK_TABLE","0x0003",       "Wire protocol: batch track table"],
            ["MAX_DETECTIONS",    "256",          "Array size cap for detection messages"],
            ["MAX_TRACKS",        "64",           "Array size cap for track table messages"],
            ["IMM_NUM_MODELS",    "5",            "Number of IMM motion models"],
        ],
        col_widths=[2.0, 1.2, 3.4]
    )

    add_heading(doc, "matrix_ops.h", 3)
    add_para(doc, "Header-only matrix utility library. All matrices are fixed-size arrays. Key functions:")
    add_table(doc,
        ["Function",               "Description"],
        [
            ["mat9_multiply(A, B)",    "9×9 matrix multiplication"],
            ["mat9_add(A, B)",         "9×9 matrix addition"],
            ["mat9_transpose(A)",      "9×9 matrix transpose"],
            ["mat9_inverse(A)",        "9×9 matrix inversion via Gauss-Jordan elimination"],
            ["mat3_inverse(A)",        "3×3 matrix inversion"],
            ["mat_9x3_multiply(A, B)", "9×3 × 3×3 mixed-dimension multiplication"],
            ["mat_3x9_multiply(A, B)", "3×9 × 9×9 mixed-dimension multiplication"],
            ["mahalanobis_dist_sq(z, z_pred, S)", "Mahalanobis distance squared: (z-z_pred)ᵀ S⁻¹ (z-z_pred)"],
        ],
        col_widths=[2.8, 3.8]
    )

    add_heading(doc, "8.2 Tracker Pipeline  (cuas_pipeline)", 2)
    add_para(doc, (
        "TrackerPipeline is the top-level orchestrator. It:"
    ))
    for item in [
        "Constructs and owns all sub-system objects (receiver, track_manager, sender)",
        "Registers the detection callback with DetectionReceiver",
        "Starts the receiver thread and an internal processor thread",
        "On each cycle: dequeues available scans, calls TrackManager::process(), calls TrackSender::send()",
        "Maintains statistics counters (received, processed, sent) and logs them periodically",
        "On shutdown: signals threads to stop, joins them, flushes the logger",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8.3 Track Manager  (cuas_track_management)", 2)
    add_para(doc, "TrackManager::process() executes the following ordered steps each call:")
    for i, step in enumerate([
        "Call Preprocessor::filter() → reduced detection list",
        "Call ClusterEngine::cluster() → cluster list",
        "For each existing track: call IMMFilter::predict() to advance state to current time",
        "Call AssociationEngine::associate() → (track_id → cluster_idx) assignment map",
        "For each assigned pair: call IMMFilter::update() with cluster measurement",
        "Record hits/misses on all tracks; decay/boost quality",
        "Call TrackInitiator::process() with unassigned clusters → new tentative tracks",
        "Apply status transitions (Tentative→Confirmed, Confirmed→Coasting, Coasting→Deleted)",
        "Apply quality and age deletion rules",
        "Classify all active tracks",
    ], 1):
        add_bullet(doc, f"Step {i}: {step}")

    page_break(doc)

    # =========================================================================
    # 9. SIMULATOR TOOLS
    # =========================================================================
    add_heading(doc, "9. Simulator and Utility Tools", 1)

    add_heading(doc, "9.1 DSP Injector  (dsp_injector)", 2)
    add_para(doc, (
        "A synthetic radar target generator that simulates a configurable number of UAS targets "
        "and transmits SPDetectionMessage packets to the tracker over UDP. Each target follows a "
        "physics-based trajectory with configurable speed, heading, climb rate, and turn rate."
    ))
    add_heading(doc, "Usage:", 3)
    add_code_block(doc, "dsp_injector [ip] [port] [num_targets] [duration_sec] [rate_ms]\n"
                        "  ip           Tracker IP address          (default: 127.0.0.1)\n"
                        "  port         Tracker listen port          (default: 5000)\n"
                        "  num_targets  Number of simulated targets  (default: 3)\n"
                        "  duration_sec Run duration in seconds      (default: 120)\n"
                        "  rate_ms      Scan interval in ms          (default: 100)")
    add_heading(doc, "Physics model:", 3)
    for item in [
        "Position updated each cycle using velocity (Cartesian then converted to spherical)",
        "Target boundary constraints: altitude 10–3000 m, range 30–20,000 m",
        "Detection probability Pd = f(range, RCS): Pd decreases with range² and increases with RCS",
        "Range noise: Gaussian σ ≈ 5 m; angle noise: Gaussian σ ≈ 0.01 rad",
        "False alarm (clutter) detections generated at a configurable rate",
        "Path loss model for strength: −30 + RCS_dBsm − 40·log₁₀(range_m)",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "9.2 Console Display Module  (display_module)", 2)
    add_para(doc, (
        "A lightweight terminal application that listens for TrackTableMessage packets and "
        "renders a continuously updating track table using ANSI terminal escape codes."
    ))
    add_code_block(doc, "display_module [listen_port]   (default port: 5001)")
    add_para(doc, "Displays columns: Track ID, Status, Classification, Range, Azimuth, Elevation, Range-rate, X, Y, Z, Quality, Hits, Misses, Age. Also prints summary counts of Confirmed / Tentative / Coasting tracks.")

    add_heading(doc, "9.3 Log Extractor  (log_extractor)", 2)
    add_para(doc, "Multi-mode binary log analysis tool.")
    add_code_block(doc,
        "log_extractor <file> <mode> [options]\n\n"
        "Modes:\n"
        "  extract           Human-readable dump of all log records with type breakdown\n"
        "  replay [speed]    Re-inject RawDetection records to tracker at configurable speed\n"
        "                    speed: 1.0 = real-time, 2.0 = double speed, 0.5 = half speed\n"
        "  csv [output.csv]  Export TrackSent records to CSV with full state vectors"
    )
    add_heading(doc, "Log record types recorded:", 3)
    add_table(doc,
        ["Record Type",    "Pipeline Stage",        "Payload Content"],
        [
            ["RawDetection",   "Receiver output",       "Full SPDetectionMessage (all detections)"],
            ["Preprocessed",   "Preprocessor output",   "Filtered detection list"],
            ["Clustered",      "Clusterer output",       "Cluster centroid list"],
            ["Predicted",      "IMM predict step",       "Per-track predicted state + covariance"],
            ["Associated",     "Association output",     "Track-cluster assignment map"],
            ["TrackInitiated", "Initiator output",       "New track initial state"],
            ["TrackUpdated",   "IMM update step",        "Updated track state + covariance"],
            ["TrackDeleted",   "Track manager deletion", "Final state at deletion"],
            ["TrackSent",      "Sender output",          "TrackTableMessage as transmitted"],
        ],
        col_widths=[1.6, 1.8, 3.2]
    )

    add_heading(doc, "9.4 Qt Display Module  (DisplayModule)", 2)
    add_para(doc, (
        "A Qt5 Widgets GUI application providing a graphical track table. It connects to the "
        "tracker's output port via a QUdpSocket and updates the display each time a "
        "TrackTableMessage arrives."
    ))
    add_heading(doc, "GUI elements:", 3)
    for item in [
        "Port spinbox (1024–65535) — configures which UDP port to listen on",
        "Start / Stop button — connects or disconnects from the UDP port",
        "Track table (14 columns): ID, Status, Classification, Range, Azimuth, Elevation, Range-rate, X, Y, Z, Quality, Hits, Misses, Age",
        "Status bar: total message count, count of Confirmed / Tentative / Coasting tracks",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Build (Qt):", 3)
    add_code_block(doc,
        "cd qt_display_module\n"
        "qmake DisplayModule.pro\n"
        "nmake   # Windows (MSVC)\n"
        "# or\n"
        "make    # Linux / macOS"
    )

    page_break(doc)

    # =========================================================================
    # 10. THREADING MODEL
    # =========================================================================
    add_heading(doc, "10. Threading Model", 1)
    add_table(doc,
        ["Thread",          "Owner",             "Responsibility"],
        [
            ["Receiver Thread", "DetectionReceiver", "Blocks on recvfrom(); deserialises; pushes to queue; signals processor"],
            ["Processor Thread","TrackerPipeline",   "Waits on condition variable; dequeues scan; runs full pipeline; sends tracks"],
            ["Main Thread",     "main.cpp",          "Initialises components; waits for SIGINT/SIGTERM; triggers graceful shutdown"],
        ],
        col_widths=[1.6, 1.8, 3.2]
    )
    add_para(doc, (
        "The queue between receiver and processor is protected by a std::mutex with a "
        "std::condition_variable. The receiver pushes and notifies; the processor waits and pops. "
        "An std::atomic<bool> running flag controls the shutdown sequence: it is set to false by "
        "the signal handler, which causes both threads to exit their loops and be join()ed by the "
        "main thread."
    ))

    # =========================================================================
    # 11. LOGGING
    # =========================================================================
    add_heading(doc, "11. Logging Subsystem", 1)

    add_heading(doc, "11.1 Binary Logger", 2)
    add_para(doc, (
        "BinaryLogger writes structured records to a binary file. Each record consists of: "
        "a fixed 4-byte magic number (0xDEADBEEF), a 1-byte record type (LogRecordType enum), "
        "an 8-byte millisecond timestamp, a 4-byte payload length, and the variable-length payload. "
        "This format enables fast random-access parsing by the log_extractor tool."
    ))

    add_heading(doc, "11.2 Console Logger", 2)
    add_para(doc, (
        "ConsoleLogger writes timestamped, level-tagged messages to stdout. It is mutex-protected "
        "for safe use from multiple threads. Log levels: DEBUG < INFO < WARN < ERROR. "
        "Messages below the configured log_level are suppressed."
    ))

    page_break(doc)

    # =========================================================================
    # 12. NETWORK PROTOCOL
    # =========================================================================
    add_heading(doc, "12. Network Protocol Details", 1)

    add_heading(doc, "12.1 Serialisation", 2)
    add_para(doc, (
        "All messages are serialised to little-endian binary using memcpy of the struct layout "
        "directly into the UDP payload buffer. No padding bytes are inserted between struct fields "
        "(the structs are designed to be naturally aligned). This means the receiver can deserialise "
        "by reading the buffer back into the struct via memcpy, provided both ends use the same ABI."
    ))

    add_heading(doc, "12.2 Default Port Assignments", 2)
    add_table(doc,
        ["Port",  "Direction",          "Protocol",  "Message Type"],
        [
            ["5000", "DSP → Tracker",      "UDP",       "SPDetectionMessage (radar detections)"],
            ["5001", "Tracker → Display",  "UDP",       "TrackTableMessage (track estimates)"],
        ],
        col_widths=[0.8, 2.0, 1.0, 2.8]
    )

    add_heading(doc, "12.3 Message ID Summary", 2)
    add_table(doc,
        ["Message ID", "Hex",    "Type"],
        [
            ["1",          "0x0001", "SPDetectionMessage"],
            ["2",          "0x0002", "TrackUpdateMessage"],
            ["3",          "0x0003", "TrackTableMessage"],
        ],
        col_widths=[1.4, 1.2, 4.0]
    )

    # =========================================================================
    # 13. RUNNING THE SYSTEM
    # =========================================================================
    add_heading(doc, "13. Running the System", 1)

    add_heading(doc, "13.1 Quick Start (Simulation Mode)", 2)
    add_para(doc, "Open three terminals from the build/Debug (or build/Release) directory:")
    add_para(doc, "Terminal 1 — Start the tracker:")
    add_code_block(doc, "cuas_tracker")
    add_para(doc, "Terminal 2 — Start the DSP injector (3 simulated targets, 120 s):")
    add_code_block(doc, "dsp_injector 127.0.0.1 5000 3 120 100")
    add_para(doc, "Terminal 3 — Start the console display:")
    add_code_block(doc, "display_module 5001")
    add_para(doc, (
        "Alternatively, launch the Qt GUI display: run DisplayModule.exe, enter port 5001, "
        "and click Start."
    ))

    add_heading(doc, "13.2 Log Replay", 2)
    add_para(doc, "After a live run, replay the captured log at double speed:")
    add_code_block(doc, "log_extractor tracker_log.bin replay 127.0.0.1 5000 2.0")
    add_para(doc, "Export all track outputs to CSV:")
    add_code_block(doc, "log_extractor tracker_log.bin csv tracks_output.csv")

    add_heading(doc, "13.3 Configuration Tuning Tips", 2)
    for tip in [
        "Increase max_coasting_cycles if tracks are being dropped during temporary occlusions.",
        "Increase gate_threshold (chi-squared) in high-clutter environments to maintain association at the cost of more false associations.",
        "Switch clustering to DBSCAN for spatially distributed swarm scenarios; use RangeBased for dense linear-flight targets.",
        "Switch association to JPDA in environments with very dense false alarms.",
        "Reduce cycle_period_ms (e.g., to 50 ms) if the radar operates at 20 Hz scan rate.",
        "Increase m_of_n_m / m_of_n_n for stricter track confirmation, reducing false tracks at the cost of confirmation latency.",
    ]:
        add_bullet(doc, tip)

    page_break(doc)

    # =========================================================================
    # 14. GLOSSARY
    # =========================================================================
    add_heading(doc, "14. Glossary", 1)
    add_table(doc,
        ["Term",        "Definition"],
        [
            ["UAS",         "Unmanned Aerial System (drone)"],
            ["Counter-UAS", "System designed to detect, track, and optionally defeat UAS threats"],
            ["DSP",         "Digital Signal Processor — the radar front-end that produces raw detections"],
            ["IMM",         "Interacting Multiple Model filter — probabilistic fusion of multiple Kalman filters"],
            ["CV",          "Constant Velocity motion model"],
            ["CA",          "Constant Acceleration motion model"],
            ["CTR",         "Coordinated Turn Rate motion model"],
            ["DBSCAN",      "Density-Based Spatial Clustering of Applications with Noise"],
            ["GNN",         "Global Nearest Neighbour data association (also called MHT-lite or Hungarian)"],
            ["JPDA",        "Joint Probabilistic Data Association — soft, weight-based association for dense environments"],
            ["M-of-N",      "Track initiation logic: M hits within N scans required to confirm a candidate"],
            ["ENU",         "East-North-Up Cartesian coordinate frame"],
            ["SNR",         "Signal-to-Noise Ratio"],
            ["RCS",         "Radar Cross-Section — a measure of target radar reflectivity"],
            ["dBsm",        "Decibels relative to one square metre — unit for RCS"],
            ["Mahalanobis", "Statistical distance accounting for covariance structure"],
            ["Coasting",    "Track state where the filter predicts forward without a matched measurement"],
            ["Quality",     "Scalar track health metric [0,1]: decays on misses, grows on hits"],
        ],
        col_widths=[1.4, 5.2]
    )

    # =========================================================================
    # Save
    # =========================================================================
    out_path = r"d:\Zoppler Projects\Tracker_CUxS\CounterUAS_Radar_Tracker_Documentation.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")

if __name__ == "__main__":
    build()
